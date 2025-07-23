import os
import json
from datetime import datetime
from langchain.tools import StructuredTool, Tool
try:
    from pydantic.v1 import BaseModel, Field  # For newer pydantic versions
except ImportError:
    from pydantic import BaseModel, Field  # For older pydantic versions
from typing import Type, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class SavePlanInput(BaseModel):
    plan: str = Field(description="The travel plan/itinerary to save (can be JSON string or regular text)")
    location: str = Field(description="The destination/location name")

def create_word_document(plan_dict: dict, location: str, filename: str) -> str:
    """Create a beautifully formatted Word document"""
    try:
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Travel Itinerary: {location.title()}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        date_para = doc.add_paragraph()
        date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y")}').italic = True
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a line break
        doc.add_paragraph()
        
        # Destination Overview Section
        doc.add_heading('Destination Overview', level=1)
        overview_table = doc.add_table(rows=0, cols=2)
        
        # Add overview details
        if isinstance(plan_dict, dict):
            if 'destination' in plan_dict:
                row = overview_table.add_row()
                row.cells[0].text = 'Location'
                row.cells[1].text = plan_dict.get('destination', location).title()
            
            if 'duration' in plan_dict:
                row = overview_table.add_row()
                row.cells[0].text = 'Duration'
                row.cells[1].text = plan_dict['duration']
            
            if 'budget_range' in plan_dict:
                row = overview_table.add_row()
                row.cells[0].text = 'Budget Level'
                row.cells[1].text = plan_dict['budget_range']
            
            if 'best_time_to_visit' in plan_dict:
                row = overview_table.add_row()
                row.cells[0].text = 'Best Time to Visit'
                row.cells[1].text = plan_dict['best_time_to_visit']
        
        doc.add_paragraph()
        
        # Trip Highlights Section
        if isinstance(plan_dict, dict) and 'highlights' in plan_dict:
            doc.add_heading('Trip Highlights', level=1)
            if isinstance(plan_dict['highlights'], list):
                for highlight in plan_dict['highlights']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(highlight)
            elif isinstance(plan_dict['highlights'], str):
                highlights = [h.strip() for h in plan_dict['highlights'].split(',')]
                for highlight in highlights:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(highlight)
        
        # Detailed Itinerary Section
        doc.add_heading('Detailed Itinerary', level=1)
        
        if isinstance(plan_dict, dict) and 'itinerary' in plan_dict:
            if isinstance(plan_dict['itinerary'], list):
                # Structured itinerary
                for day in plan_dict['itinerary']:
                    if isinstance(day, dict):
                        # Day heading
                        day_title = f"Day {day.get('day_number', '?')}: {day.get('title', 'Exploration Day')}"
                        doc.add_heading(day_title, level=2)
                        
                        # Activities
                        if 'activities' in day and isinstance(day['activities'], list):
                            for activity in day['activities']:
                                if isinstance(activity, dict):
                                    # Time and activity
                                    activity_para = doc.add_paragraph()
                                    time_run = activity_para.add_run(f"{activity.get('time', 'Time')}: ")
                                    time_run.bold = True
                                    activity_para.add_run(activity.get('activity', ''))
                                    
                                    # Location
                                    if activity.get('location'):
                                        location_para = doc.add_paragraph()
                                        location_run = location_para.add_run("Location: ")
                                        location_run.bold = True
                                        location_para.add_run(activity['location'])
                                    
                                    # Cost
                                    if activity.get('cost'):
                                        cost_para = doc.add_paragraph()
                                        cost_run = cost_para.add_run("Cost: ")
                                        cost_run.bold = True
                                        cost_para.add_run(activity['cost'])
                                    
                                    # Tips
                                    if activity.get('tips'):
                                        tips_para = doc.add_paragraph()
                                        tips_run = tips_para.add_run("Tips: ")
                                        tips_run.bold = True
                                        tips_para.add_run(activity['tips'])
                                    
                                    doc.add_paragraph()  # Add spacing
                        
                        # Meals section
                        if 'meals' in day and day['meals']:
                            meals_para = doc.add_paragraph()
                            meals_para.add_run('Recommended Restaurants:').bold = True
                            for meal in day['meals']:
                                doc.add_paragraph(f'• {meal}', style='List Bullet')
                        
                        # Accommodation
                        if 'accommodation' in day and day['accommodation']:
                            acc_para = doc.add_paragraph()
                            acc_para.add_run('Accommodation: ').bold = True
                            acc_para.add_run(day['accommodation'])
                        
                        doc.add_paragraph()  # Add spacing between days
                        
            elif isinstance(plan_dict['itinerary'], str):
                # Text-based itinerary
                lines = plan_dict['itinerary'].split('\n')
                for line in lines:
                    if line.strip():
                        if line.lower().startswith('day'):
                            doc.add_heading(line, level=2)
                        else:
                            doc.add_paragraph(line)
        
        # Packing Essentials Section
        if isinstance(plan_dict, dict) and 'packing_tips' in plan_dict and plan_dict['packing_tips']:
            doc.add_heading('Packing Essentials', level=1)
            if isinstance(plan_dict['packing_tips'], list):
                for item in plan_dict['packing_tips']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(item)
        
        # Local Tips Section
        if isinstance(plan_dict, dict) and 'local_tips' in plan_dict and plan_dict['local_tips']:
            doc.add_heading('Local Tips & Cultural Notes', level=1)
            if isinstance(plan_dict['local_tips'], list):
                for tip in plan_dict['local_tips']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(tip)
        
        # Emergency Contacts Section
        if isinstance(plan_dict, dict) and 'emergency_contacts' in plan_dict and plan_dict['emergency_contacts']:
            doc.add_heading('Emergency Contacts', level=1)
            if isinstance(plan_dict['emergency_contacts'], list):
                for contact in plan_dict['emergency_contacts']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(contact)
        
        # Important Information Section
        doc.add_page_break()
        doc.add_heading('Important Information', level=1)
        
        # Before You Go Checklist
        doc.add_heading('Before You Go Checklist', level=2)
        checklist = [
            'Check passport validity (6+ months)',
            'Verify visa requirements',
            'Get travel insurance',
            'Inform bank of travel plans',
            'Make copies of important documents',
            'Download offline maps',
            'Check vaccination requirements'
        ]
        
        for item in checklist:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f'☐ {item}')
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run('This travel plan was generated by your AI Travel Assistant').italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer2 = doc.add_paragraph()
        footer2.add_run('Have an amazing trip!').bold = True
        footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save the document
        doc.save(filename)
        return filename
        
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return None

def save_plan(plan: str, location: str) -> str:
    """Save the travel plan as a Word document and markdown file"""
    try:
        os.makedirs("plans", exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Parse the plan
        try:
            if isinstance(plan, str) and (plan.strip().startswith('{') or plan.strip().startswith('[')):
                plan_dict = json.loads(plan)
            else:
                plan_dict = {"raw_plan": plan}
        except Exception as e:
            print(f"Error parsing plan JSON: {e}")
            plan_dict = {"raw_plan": plan}
        
        # Clean location name for filename
        clean_location = location.replace(' ', '_').replace(',', '').replace('-', '_').lower()
        
        # Create Word document
        word_filename = f"plans/{clean_location}_{timestamp}.docx"
        result = create_word_document(plan_dict, location, word_filename)
        
        if not result:
            return "Error creating Word document"
        
        # Also create markdown as backup
        md_filename = f"plans/{clean_location}_{timestamp}.md"
        
        # Create markdown content
        markdown_content = f"""# Travel Itinerary: {location.title()}

**Generated on:** {datetime.now().strftime("%B %d, %Y at %I:%M %p")}

---

## Destination Overview

**Location:** {location.title()}  
"""
        
        if isinstance(plan_dict, dict):
            if 'duration' in plan_dict:
                markdown_content += f"**Duration:** {plan_dict['duration']}  \n"
            if 'budget_range' in plan_dict:
                markdown_content += f"**Budget:** {plan_dict['budget_range']}  \n"
            
            markdown_content += "\n"
            
            # Add highlights
            if 'highlights' in plan_dict and plan_dict['highlights']:
                markdown_content += "## Trip Highlights\n\n"
                if isinstance(plan_dict['highlights'], list):
                    for highlight in plan_dict['highlights']:
                        markdown_content += f"- {highlight}\n"
                markdown_content += "\n"
            
            # Add itinerary
            if 'itinerary' in plan_dict and plan_dict['itinerary']:
                markdown_content += "## Detailed Itinerary\n\n"
                if isinstance(plan_dict['itinerary'], list):
                    for day in plan_dict['itinerary']:
                        if isinstance(day, dict):
                            markdown_content += f"### Day {day.get('day_number', '?')}: {day.get('title', 'Exploration Day')}\n\n"
                            
                            if 'activities' in day and day['activities']:
                                for activity in day['activities']:
                                    if isinstance(activity, dict):
                                        markdown_content += f"**{activity.get('time', 'Time')}:** {activity.get('activity', '')}\n"
                                        if activity.get('location'):
                                            markdown_content += f"- *Location:* {activity['location']}\n"
                                        if activity.get('cost'):
                                            markdown_content += f"- *Cost:* {activity['cost']}\n"
                                        if activity.get('tips'):
                                            markdown_content += f"- *Tip:* {activity['tips']}\n"
                                        markdown_content += "\n"
                            
                            if 'meals' in day and day['meals']:
                                markdown_content += "**Recommended Restaurants:**\n"
                                for meal in day['meals']:
                                    markdown_content += f"- {meal}\n"
                                markdown_content += "\n"
            
            # Add other sections
            if 'local_tips' in plan_dict and plan_dict['local_tips']:
                markdown_content += "## Local Tips & Cultural Notes\n\n"
                if isinstance(plan_dict['local_tips'], list):
                    for tip in plan_dict['local_tips']:
                        markdown_content += f"- {tip}\n"
                markdown_content += "\n"
            
            if 'packing_tips' in plan_dict and plan_dict['packing_tips']:
                markdown_content += "## Packing Essentials\n\n"
                if isinstance(plan_dict['packing_tips'], list):
                    for item in plan_dict['packing_tips']:
                        markdown_content += f"- {item}\n"
                markdown_content += "\n"
        
        with open(md_filename, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        
        # Save JSON backup
        json_filename = md_filename.replace('.md', '.json')
        with open(json_filename, "w", encoding="utf-8") as f:
            if isinstance(plan_dict, dict) and 'raw_plan' not in plan_dict:
                json.dump(plan_dict, f, indent=2)
            else:
                json.dump({"location": location, "plan": plan}, f, indent=2)
        
        # Get just the filename for return
        word_file_name = os.path.basename(word_filename)
        md_file_name = os.path.basename(md_filename)
        
        return f"✅ Plan saved successfully!\n📄 Word Document: {word_file_name}\n📝 Markdown: {md_file_name}"
    except Exception as e:
        print(f"Error in save_plan: {e}")
        return f"Error saving plan: {str(e)}"

def suggest_trip(city: str) -> str:
    """Suggest things to do in a given location"""
    suggestions = {
        "bali": [
            "Visit ancient temples like Tanah Lot and Uluwatu",
            "Explore rice terraces in Ubud",
            "Relax on beautiful beaches in Seminyak",
            "Experience traditional Balinese culture",
            "Try local cuisine and cooking classes"
        ],
        "paris": [
            "Visit the Eiffel Tower and Trocadéro Gardens",
            "Explore the Louvre Museum and see the Mona Lisa",
            "Stroll through Montmartre and visit Sacré-Cœur",
            "Take a Seine River cruise",
            "Experience local cuisine in Le Marais district"
        ],
        "munnar": [
            "Visit tea plantations and tea museums",
            "Go trekking in the Western Ghats",
            "See wildlife at Eravikulam National Park",
            "Enjoy boating at Mattupetty Dam",
            "Experience local spice gardens"
        ],
        "goa": [
            "Relax on pristine beaches like Anjuna and Calangute",
            "Explore Portuguese colonial architecture",
            "Visit spice plantations",
            "Experience vibrant nightlife",
            "Try water sports and beach activities"
        ],
        "japan": [
            "Visit traditional temples and shrines",
            "Experience cherry blossom season",
            "Try authentic Japanese cuisine",
            "Explore modern cities like Tokyo and Osaka",
            "Experience traditional ryokan stays"
        ],
        "kerala": [
            "Cruise through the backwaters of Alleppey",
            "Visit tea plantations in Munnar",
            "Experience Ayurvedic treatments",
            "Explore wildlife in Periyar National Park",
            "Watch traditional Kathakali performances"
        ],
        "tokyo": [
            "Visit traditional temples like Senso-ji",
            "Experience the bustling Shibuya crossing",
            "Try authentic sushi at Tsukiji market",
            "Explore modern districts like Harajuku",
            "Take day trips to Mount Fuji"
        ]
    }
    
    city_lower = city.lower()
    
    # Check for partial matches
    for key in suggestions.keys():
        if key in city_lower or city_lower in key:
            return f"Top 5 things to do in {city}:\n" + "\n".join([f"{i+1}. {item}" for i, item in enumerate(suggestions[key])])
    
    # Default suggestions if city not found
    return f"""Top 5 things to do in {city}:
1. Visit historical sites and landmarks
2. Try traditional local cuisine and food tours
3. Explore museums, art galleries, or cultural centers
4. Go sightseeing or take walking tours
5. Experience local markets and shopping districts"""

# Create tools list
tools = [
    StructuredTool.from_function(
        func=save_plan,
        name="save_plan",
        description="Save the generated itinerary to a file. Requires both plan content and location name.",
        args_schema=SavePlanInput
    ),
    Tool(
        name="suggest_trip",
        func=suggest_trip,
        description="Suggest activities and attractions for a city or destination"
    )
]